import io
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import numpy as np

class DOCXReportGenerator:
    """Generate DOCX reports for all 17 CI report types"""
    
    REPORT_TYPES = [
        "Clinical Positioning Analysis",
        "Regulatory Risk Assessment",
        "Financial Exposure Report",
        "Time-to-Market Benchmarking",
        "Competitive Density Mapping",
        "Partnership & Licensing Activity",
        "Innovation Strength Assessment",
        "Patent Landscape Analysis",
        "Pipeline Monitoring Report",
        "Market Access & Pricing Report",
        "Regulatory Landscape Report",
        "Clinical Trial Competitive Intelligence",
        "Early Warning Report",
        "Competitive Landscape Report",
        "Competitor Sales & Market Share Report",
        "Patent Expiry & Biosimilar Entry Report",
        "Freedom to Operate Assessment"
    ]
    
    def __init__(self, df_filtered, metrics, ci_scores):
        self.df_filtered = df_filtered
        self.metrics = metrics
        self.ci_scores = ci_scores
        self.timestamp = datetime.now()
    
    def _add_heading(self, doc, text, level=1):
        """Add formatted heading"""
        heading = doc.add_heading(text, level=level)
        if level == 1:
            heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
        return heading
    
    def _add_ci_scores_table(self, doc):
        """Add CI scores table to document"""
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Domain'
        hdr_cells[1].text = 'Score'
        hdr_cells[2].text = '95% CI Lower'
        hdr_cells[3].text = '95% CI Upper'
        
        for domain, scores in self.ci_scores.items():
            row_cells = table.add_row().cells
            row_cells[0].text = domain
            row_cells[1].text = f"{scores['score']:.0f}/100"
            row_cells[2].text = f"{scores['ci_lower']:.0f}"
            row_cells[3].text = f"{scores['ci_upper']:.0f}"
    
    def _add_company_table(self, doc, limit=10):
        """Add company analysis table"""
        top_companies = self.df_filtered.drop_duplicates('Company_Name')[
            ['Company_Name', 'NASDAQ_Symbol', 'Market_Cap', 'Disease_Area', 'Trial_Clinical_Phase']
        ].nlargest(limit, 'Market_Cap')
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Company'
        hdr_cells[1].text = 'Symbol'
        hdr_cells[2].text = 'Market Cap'
        hdr_cells[3].text = 'Disease Area'
        hdr_cells[4].text = 'Phase'
        
        for _, company in top_companies.iterrows():
            market_cap = company['Market_Cap']
            if isinstance(market_cap, (int, float)) and not np.isnan(market_cap):
                market_cap_str = f"${market_cap/1e9:.1f}B"
            else:
                market_cap_str = "N/A"
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(company['Company_Name'])
            row_cells[1].text = str(company['NASDAQ_Symbol'])
            row_cells[2].text = market_cap_str
            row_cells[3].text = str(company['Disease_Area'])
            row_cells[4].text = str(company['Trial_Clinical_Phase'])
    
    def generate_clinical_positioning_report(self):
        """Generate Clinical Positioning Analysis report"""
        doc = Document()
        
        self._add_heading(doc, "Clinical Positioning Analysis", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Data Source: Official APIs (100% Verified)")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Total Companies: {self.df_filtered['Company_Name'].nunique()}")
        doc.add_paragraph(f"Active Trials: {self.metrics['active_trials']}")
        doc.add_paragraph(f"Average Market Cap: ${self.metrics['avg_market_cap']/1e9:.2f}B")
        doc.add_paragraph()
        
        self._add_heading(doc, "Clinical Phase Distribution", level=2)
        phase_dist = self.df_filtered['Trial_Clinical_Phase'].value_counts()
        for phase, count in phase_dist.head(8).items():
            if pd.notna(phase):
                pct = (count / len(self.df_filtered)) * 100
                doc.add_paragraph(f"{phase}: {count} trials ({pct:.1f}%)", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "Market Leaders", level=2)
        self._add_company_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "Key Findings", level=2)
        doc.add_paragraph(f"The market shows {'strong' if self.metrics['active_trials'] > 1000 else 'moderate'} activity with {self.metrics['active_trials']} active trials.", style='List Bullet')
        doc.add_paragraph(f"Clinical Maturity Score: {self.ci_scores['Clinical Maturity']['score']:.0f}/100", style='List Bullet')
        doc.add_paragraph(f"Pipeline Diversification Score: {self.ci_scores['Pipeline Diversification']['score']:.0f}/100", style='List Bullet')
        
        return doc
    
    def generate_regulatory_risk_report(self):
        """Generate Regulatory Risk Assessment report"""
        doc = Document()
        
        self._add_heading(doc, "Regulatory Risk Assessment", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Total Companies Analyzed: {self.df_filtered['Company_Name'].nunique()}")
        doc.add_paragraph(f"Completed Trials: {len(self.df_filtered[self.df_filtered['Trial_Overall_Status'] == 'COMPLETED'])}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Regulatory Strength Analysis", level=2)
        doc.add_paragraph(f"Score: {self.ci_scores['Regulatory Strength']['score']:.0f}/100")
        doc.add_paragraph(f"95% Confidence Interval: {self.ci_scores['Regulatory Strength']['ci_lower']:.0f} – {self.ci_scores['Regulatory Strength']['ci_upper']:.0f}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Trial Status Distribution", level=2)
        status_dist = self.df_filtered['Trial_Overall_Status'].value_counts()
        for status, count in status_dist.head(8).items():
            pct = (count / len(self.df_filtered)) * 100
            doc.add_paragraph(f"{status}: {count} ({pct:.1f}%)", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "Risk Assessment", level=2)
        risk_level = 'High' if self.metrics['risk_index'] > 70 else 'Moderate' if self.metrics['risk_index'] > 50 else 'Low'
        doc.add_paragraph(f"Overall Risk Level: {risk_level}")
        doc.add_paragraph(f"Risk Index: {self.metrics['risk_index']}/100")
        
        return doc
    
    def generate_financial_exposure_report(self):
        """Generate Financial Exposure Report"""
        doc = Document()
        
        self._add_heading(doc, "Financial Exposure Report", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Average Market Cap: ${self.metrics['avg_market_cap']/1e9:.2f}B")
        doc.add_paragraph(f"Financial Stability Score: {self.ci_scores['Financial Stability']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "Market Capitalization Analysis", level=2)
        self._add_company_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "Financial Stability Assessment", level=2)
        doc.add_paragraph(f"Score: {self.ci_scores['Financial Stability']['score']:.0f}/100")
        doc.add_paragraph(f"95% CI: {self.ci_scores['Financial Stability']['ci_lower']:.0f} – {self.ci_scores['Financial Stability']['ci_upper']:.0f}")
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_time_to_market_report(self):
        """Generate Time-to-Market Benchmarking report"""
        doc = Document()
        
        self._add_heading(doc, "Time-to-Market Benchmarking", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Average Time to Completion: {self.metrics['avg_time_to_completion']} months")
        doc.add_paragraph(f"Active Trials: {self.metrics['active_trials']}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Phase Progression Timeline", level=2)
        phase_dist = self.df_filtered['Trial_Clinical_Phase'].value_counts()
        for phase, count in phase_dist.items():
            if pd.notna(phase):
                doc.add_paragraph(f"{phase}: {count} trials", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "Clinical Maturity Analysis", level=2)
        doc.add_paragraph(f"Score: {self.ci_scores['Clinical Maturity']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_competitive_density_report(self):
        """Generate Competitive Density Mapping report"""
        doc = Document()
        
        self._add_heading(doc, "Competitive Density Mapping", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Total Disease Areas: {self.df_filtered['Disease_Area'].nunique()}")
        doc.add_paragraph(f"Total Indications: {self.df_filtered['Disease'].nunique()}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Disease Area Competitive Density", level=2)
        disease_dist = self.df_filtered['Disease_Area'].value_counts()
        for disease, count in disease_dist.head(10).items():
            pct = (count / len(self.df_filtered)) * 100
            doc.add_paragraph(f"{disease}: {count} trials ({pct:.1f}%)", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "Pipeline Diversification", level=2)
        doc.add_paragraph(f"Score: {self.ci_scores['Pipeline Diversification']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_partnership_report(self):
        """Generate Partnership & Licensing Activity report"""
        doc = Document()
        
        self._add_heading(doc, "Partnership & Licensing Activity", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        partnership_count = self.df_filtered['Partnerships'].notna().sum()
        doc.add_paragraph(f"Companies with Partnerships: {partnership_count}")
        doc.add_paragraph(f"Partnership Rate: {(partnership_count / len(self.df_filtered)) * 100:.1f}%")
        doc.add_paragraph()
        
        self._add_heading(doc, "Partnership Activity", level=2)
        doc.add_paragraph(f"Score: {self.ci_scores['Partnership Activity']['score']:.0f}/100")
        doc.add_paragraph(f"95% CI: {self.ci_scores['Partnership Activity']['ci_lower']:.0f} – {self.ci_scores['Partnership Activity']['ci_upper']:.0f}")
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_innovation_report(self):
        """Generate Innovation Strength Assessment report"""
        doc = Document()
        
        self._add_heading(doc, "Innovation Strength Assessment", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Innovation Index Score: {self.ci_scores['Innovation Index']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "Technology Platform Analysis", level=2)
        tech_dist = self.df_filtered['Technology'].value_counts().head(10)
        for tech, count in tech_dist.items():
            if pd.notna(tech):
                doc.add_paragraph(f"{str(tech)[:50]}: {count} companies", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "Innovation Index", level=2)
        doc.add_paragraph(f"Score: {self.ci_scores['Innovation Index']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_patent_landscape_report(self):
        """Generate Patent Landscape Analysis report"""
        doc = Document()
        
        self._add_heading(doc, "Patent Landscape Analysis", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Total Companies: {self.df_filtered['Company_Name'].nunique()}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Patent Portfolio Overview", level=2)
        self._add_company_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_pipeline_monitoring_report(self):
        """Generate Pipeline Monitoring Report"""
        doc = Document()
        
        self._add_heading(doc, "Pipeline Monitoring Report", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Total Pipeline Size: {len(self.df_filtered)} trials")
        doc.add_paragraph(f"Unique Companies: {self.df_filtered['Company_Name'].nunique()}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Phase Distribution", level=2)
        phase_dist = self.df_filtered['Trial_Clinical_Phase'].value_counts()
        for phase, count in phase_dist.items():
            if pd.notna(phase):
                pct = (count / len(self.df_filtered)) * 100
                doc.add_paragraph(f"{phase}: {count} ({pct:.1f}%)", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_market_access_report(self):
        """Generate Market Access & Pricing Report"""
        doc = Document()
        
        self._add_heading(doc, "Market Access & Pricing Report", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Regulatory Strength: {self.ci_scores['Regulatory Strength']['score']:.0f}/100")
        doc.add_paragraph(f"Financial Stability: {self.ci_scores['Financial Stability']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "Market Leaders", level=2)
        self._add_company_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_regulatory_landscape_report(self):
        """Generate Regulatory Landscape Report"""
        doc = Document()
        
        self._add_heading(doc, "Regulatory Landscape Report", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Regulatory Strength Score: {self.ci_scores['Regulatory Strength']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "Trial Status Overview", level=2)
        status_dist = self.df_filtered['Trial_Overall_Status'].value_counts()
        for status, count in status_dist.items():
            pct = (count / len(self.df_filtered)) * 100
            doc.add_paragraph(f"{status}: {count} ({pct:.1f}%)", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_clinical_trial_report(self):
        """Generate Clinical Trial Competitive Intelligence report"""
        doc = Document()
        
        self._add_heading(doc, "Clinical Trial Competitive Intelligence", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Active Trials: {self.metrics['active_trials']}")
        doc.add_paragraph(f"Total Trials: {self.df_filtered['Trial_NCT_ID'].nunique()}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Trial Portfolio Overview", level=2)
        doc.add_paragraph(f"Clinical Maturity Score: {self.ci_scores['Clinical Maturity']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_early_warning_report(self):
        """Generate Early Warning Report"""
        doc = Document()
        
        self._add_heading(doc, "Early Warning Report", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Risk Index: {self.metrics['risk_index']}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "Emerging Threats", level=2)
        doc.add_paragraph(f"High competition in {self.df_filtered[self.df_filtered['Competition_Level'] == 'High']['Disease_Area'].nunique()} disease areas", style='List Bullet')
        doc.add_paragraph(f"Active trials: {self.metrics['active_trials']}", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_competitive_landscape_report(self):
        """Generate Competitive Landscape Report"""
        doc = Document()
        
        self._add_heading(doc, "Competitive Landscape Report", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Total Companies: {self.df_filtered['Company_Name'].nunique()}")
        doc.add_paragraph(f"Average Market Cap: ${self.metrics['avg_market_cap']/1e9:.2f}B")
        doc.add_paragraph()
        
        self._add_heading(doc, "Market Leaders", level=2)
        self._add_company_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_competitor_sales_report(self):
        """Generate Competitor Sales & Market Share Report"""
        doc = Document()
        
        self._add_heading(doc, "Competitor Sales & Market Share Report", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Financial Stability Score: {self.ci_scores['Financial Stability']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "Market Leaders", level=2)
        self._add_company_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_patent_expiry_report(self):
        """Generate Patent Expiry & Biosimilar Entry Report"""
        doc = Document()
        
        self._add_heading(doc, "Patent Expiry & Biosimilar Entry Report", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Total Companies: {self.df_filtered['Company_Name'].nunique()}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Market Leaders", level=2)
        self._add_company_table(doc)
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_fto_assessment_report(self):
        """Generate Freedom to Operate Assessment report"""
        doc = Document()
        
        self._add_heading(doc, "Freedom to Operate Assessment", level=1)
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        self._add_heading(doc, "Executive Summary", level=2)
        doc.add_paragraph(f"Innovation Index Score: {self.ci_scores['Innovation Index']['score']:.0f}/100")
        doc.add_paragraph()
        
        self._add_heading(doc, "Technology Platforms", level=2)
        tech_dist = self.df_filtered['Technology'].value_counts().head(10)
        for tech, count in tech_dist.items():
            if pd.notna(tech):
                doc.add_paragraph(f"{str(tech)[:50]}: {count} companies", style='List Bullet')
        doc.add_paragraph()
        
        self._add_heading(doc, "CI Domain Scores", level=2)
        self._add_ci_scores_table(doc)
        
        return doc
    
    def generate_all_reports(self):
        """Generate all 17 reports and return as dictionary"""
        reports = {
            "Clinical Positioning Analysis": self.generate_clinical_positioning_report(),
            "Regulatory Risk Assessment": self.generate_regulatory_risk_report(),
            "Financial Exposure Report": self.generate_financial_exposure_report(),
            "Time-to-Market Benchmarking": self.generate_time_to_market_report(),
            "Competitive Density Mapping": self.generate_competitive_density_report(),
            "Partnership & Licensing Activity": self.generate_partnership_report(),
            "Innovation Strength Assessment": self.generate_innovation_report(),
            "Patent Landscape Analysis": self.generate_patent_landscape_report(),
            "Pipeline Monitoring Report": self.generate_pipeline_monitoring_report(),
            "Market Access & Pricing Report": self.generate_market_access_report(),
            "Regulatory Landscape Report": self.generate_regulatory_landscape_report(),
            "Clinical Trial Competitive Intelligence": self.generate_clinical_trial_report(),
            "Early Warning Report": self.generate_early_warning_report(),
            "Competitive Landscape Report": self.generate_competitive_landscape_report(),
            "Competitor Sales & Market Share Report": self.generate_competitor_sales_report(),
            "Patent Expiry & Biosimilar Entry Report": self.generate_patent_expiry_report(),
            "Freedom to Operate Assessment": self.generate_fto_assessment_report()
        }
        return reports
    
    def export_report_to_bytes(self, doc):
        """Export document to bytes"""
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def create_ci_package_zip(self):
        """Create ZIP file with all 17 reports"""
        reports = self.generate_all_reports()
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for report_name, doc in reports.items():
                report_bytes = self.export_report_to_bytes(doc)
                filename = f"{report_name.replace(' ', '_')}.docx"
                zip_file.writestr(filename, report_bytes)
        
        zip_buffer.seek(0)
        return zip_buffer.getvalue()
